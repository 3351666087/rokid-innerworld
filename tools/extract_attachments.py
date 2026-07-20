import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
from datetime import datetime
from pathlib import Path

import pdfplumber
from docx import Document
from openpyxl import load_workbook
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
CHAT_DIR = ROOT / "群聊_Rokid"
ANALYSIS_DIR = ROOT / "analysis"
ATTACHMENTS_DIR = CHAT_DIR / "attachments"
MEDIA_DIR = CHAT_DIR / "media"
EXTRACT_DIR = ANALYSIS_DIR / "extracted_attachments"
MEDIA_EXTRACT_DIR = ANALYSIS_DIR / "media_index"
TMP_RENDER_DIR = ROOT / "tmp" / "pdfs"
CODEX_DEPS = Path(os.environ.get(
    "CODEX_RUNTIME_DEPS",
    Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies",
))
POPPLER_BIN = Path(os.environ.get("CODEX_POPPLER_BIN", CODEX_DEPS / "bin"))


def find_poppler_tool(name):
    candidates = [
        CODEX_DEPS / "native" / "poppler" / "Library" / "bin" / f"{name}.exe",
        POPPLER_BIN / f"{name}.exe",
        POPPLER_BIN / "Library" / "bin" / f"{name}.exe",
        POPPLER_BIN / f"{name}.cmd",
        CODEX_DEPS / "native" / "poppler" / "bin" / f"{name}.cmd",
        CODEX_DEPS / "bin" / f"{name}.cmd",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


PDFINFO = find_poppler_tool("pdfinfo")
PDFTOPPM = find_poppler_tool("pdftoppm")


def read_json(name):
    return json.loads((ANALYSIS_DIR / name).read_text(encoding="utf-8"))


def write_json(path, payload):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def normalize_name(name):
    stem = Path(name).stem
    suffix = Path(name).suffix.lower()
    stem = re.sub(r"\(\d+\)$", "", stem).strip()
    return f"{stem}{suffix}".lower()


def sha256_file(path):
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def safe_stem(name):
    stem = Path(name).stem
    stem = re.sub(r"[<>:\"/\\|?*\x00-\x1f]", "_", stem)
    return stem[:120]


def collect_files(root, exts=None):
    if not root.exists():
        return []
    files = []
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        if exts and path.suffix.lower() not in exts:
            continue
        files.append(path)
    return sorted(files, key=lambda p: str(p).lower())


def run_text_command(args, timeout=60):
    try:
        result = subprocess.run(
            [str(a) for a in args],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=timeout,
        )
        return {
            "ok": result.returncode == 0,
            "returncode": result.returncode,
            "stdout": result.stdout,
            "stderr": result.stderr,
        }
    except Exception as exc:
        return {"ok": False, "returncode": None, "stdout": "", "stderr": repr(exc)}


def pdf_page_count(path):
    info = run_text_command([PDFINFO, path], timeout=30) if PDFINFO else None
    if info and info["ok"]:
        match = re.search(r"^Pages:\s+(\d+)", info["stdout"], re.MULTILINE)
        if match:
            return int(match.group(1)), info
    try:
        with pdfplumber.open(path) as pdf:
            return len(pdf.pages), info
    except Exception:
        return None, info


def extract_pdf(path, out_dir):
    pages = []
    errors = []
    tables = []
    try:
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
                page_tables = []
                try:
                    raw_tables = page.extract_tables() or []
                    for table in raw_tables:
                        page_tables.append(table)
                except Exception as exc:
                    errors.append(f"page {i} table extraction: {exc!r}")
                pages.append({
                    "page": i,
                    "width": float(page.width),
                    "height": float(page.height),
                    "text": text,
                    "char_count": len(text),
                    "table_count": len(page_tables),
                })
                for t_index, table in enumerate(page_tables, start=1):
                    tables.append({"page": i, "tableIndex": t_index, "rows": table})
    except Exception as exc:
        errors.append(f"pdf extraction failed: {exc!r}")

    text_lines = []
    for page in pages:
        text_lines.append(f"\n\n===== PAGE {page['page']} =====\n")
        text_lines.append(page["text"])
    text_path = out_dir / "text.txt"
    text_path.write_text("".join(text_lines).strip() + "\n", encoding="utf-8")

    if tables:
        tables_path = out_dir / "tables.json"
        write_json(tables_path, tables)
    else:
        tables_path = None

    metadata_path = out_dir / "pages.json"
    write_json(metadata_path, {"pages": pages, "errors": errors})
    return {
        "textPath": str(text_path),
        "pagesPath": str(metadata_path),
        "tablesPath": str(tables_path) if tables_path else None,
        "pageCount": len(pages),
        "textChars": sum(p["char_count"] for p in pages),
        "tableCount": len(tables),
        "errors": errors,
    }


def extract_docx(path, out_dir):
    doc = Document(path)
    lines = []
    paragraphs = []
    for idx, para in enumerate(doc.paragraphs, start=1):
        text = para.text or ""
        if text.strip():
            paragraphs.append({"index": idx, "style": para.style.name if para.style else None, "text": text})
            lines.append(text)
    tables = []
    for t_idx, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append({"tableIndex": t_idx, "rows": rows})
        lines.append(f"\n[Table {t_idx}]")
        for row in rows:
            lines.append(" | ".join(cell.replace("\n", " / ") for cell in row))
    text_path = out_dir / "text.txt"
    text_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
    structure_path = out_dir / "docx_structure.json"
    write_json(structure_path, {"paragraphs": paragraphs, "tables": tables})
    return {
        "textPath": str(text_path),
        "structurePath": str(structure_path),
        "paragraphCount": len(paragraphs),
        "tableCount": len(tables),
        "textChars": len("\n".join(lines)),
        "errors": [],
    }


def extract_xlsx(path, out_dir):
    wb = load_workbook(path, data_only=False, read_only=False)
    data_wb = load_workbook(path, data_only=True, read_only=True)
    workbook = []
    text_lines = []
    for ws in wb.worksheets:
        data_ws = data_wb[ws.title]
        sheet = {
            "title": ws.title,
            "maxRow": ws.max_row,
            "maxColumn": ws.max_column,
            "rows": [],
            "formulas": [],
            "mergedCells": [str(rng) for rng in ws.merged_cells.ranges],
        }
        csv_path = out_dir / f"{safe_stem(ws.title)}.csv"
        with csv_path.open("w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            for r in range(1, ws.max_row + 1):
                values = []
                formulas = []
                for c in range(1, ws.max_column + 1):
                    cell = ws.cell(r, c)
                    value = cell.value
                    data_value = data_ws.cell(r, c).value
                    values.append(data_value if data_value is not None else value)
                    if isinstance(value, str) and value.startswith("="):
                        formulas.append({"cell": cell.coordinate, "formula": value, "cachedValue": data_value})
                writer.writerow(values)
                row_text = ["" if v is None else str(v) for v in values]
                if any(x.strip() for x in row_text):
                    sheet["rows"].append({"row": r, "values": row_text})
                sheet["formulas"].extend(formulas)
        workbook.append(sheet)
        text_lines.append(f"\n===== SHEET {ws.title} ({ws.max_row} x {ws.max_column}) =====")
        for row in sheet["rows"]:
            text_lines.append(f"R{row['row']}: " + " | ".join(row["values"]))
        if sheet["formulas"]:
            text_lines.append("[Formulas]")
            for formula in sheet["formulas"]:
                text_lines.append(f"{formula['cell']}: {formula['formula']} => {formula['cachedValue']}")
    text_path = out_dir / "text.txt"
    text_path.write_text("\n".join(text_lines).strip() + "\n", encoding="utf-8")
    workbook_path = out_dir / "workbook.json"
    write_json(workbook_path, workbook)
    return {
        "textPath": str(text_path),
        "workbookPath": str(workbook_path),
        "sheetCount": len(workbook),
        "textChars": len("\n".join(text_lines)),
        "formulaCount": sum(len(s["formulas"]) for s in workbook),
        "errors": [],
    }


def extract_md(path, out_dir):
    text = path.read_text(encoding="utf-8", errors="replace")
    text_path = out_dir / "text.txt"
    text_path.write_text(text, encoding="utf-8")
    return {"textPath": str(text_path), "textChars": len(text), "errors": []}


def render_pdf_samples(path, item):
    if not PDFTOPPM:
        return []
    page_count = item.get("extraction", {}).get("pageCount") or item.get("pageCount") or 0
    if not page_count:
        return []
    pages = sorted(set([1, min(2, page_count), min(3, page_count), page_count]))
    render_dir = TMP_RENDER_DIR / safe_stem(path.name)
    render_dir.mkdir(parents=True, exist_ok=True)
    rendered = []
    for page in pages:
        prefix = render_dir / f"page-{page:03d}"
        result = run_text_command([PDFTOPPM, "-f", page, "-l", page, "-png", "-r", "130", path, prefix], timeout=90)
        files = sorted(render_dir.glob(f"page-{page:03d}-*.png"))
        rendered.append({
            "page": page,
            "ok": bool(files) and result["ok"],
            "files": [str(f) for f in files],
            "stderr": result["stderr"][-1000:],
        })
    return rendered


def build_attachment_index():
    messages = read_json("messages.json")
    file_messages = []
    for message in messages:
        if message.get("type") != 4:
            continue
        match = re.match(r"^\[文件\]\s*(.+)$", message.get("content", "").strip())
        file_name = match.group(1).strip() if match else message.get("content", "").strip()
        file_messages.append({
            "messageIndex": message["index"],
            "line": message["line"],
            "time": message["time"],
            "sender": message["accountName"],
            "fileName": file_name,
            "normalizedFileName": normalize_name(file_name),
            "content": message["content"],
        })

    recovered = []
    for path in collect_files(ATTACHMENTS_DIR):
        stat = path.stat()
        recovered.append({
            "path": str(path),
            "relativePath": str(path.relative_to(ROOT)).replace("\\", "/"),
            "fileName": path.name,
            "normalizedFileName": normalize_name(path.name),
            "extension": path.suffix.lower(),
            "bytes": stat.st_size,
            "lastWriteTime": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
            "sha256": sha256_file(path),
        })

    by_norm = {}
    for item in recovered:
        by_norm.setdefault(item["normalizedFileName"], []).append(item)

    index = []
    for message in file_messages:
        candidates = by_norm.get(message["normalizedFileName"], [])
        recovered_item = candidates[0] if candidates else None
        row = {
            **message,
            "status": "matched" if recovered_item else "missing",
            "recovered": recovered_item,
        }
        index.append(row)

    matched_paths = {row["recovered"]["path"] for row in index if row["recovered"]}
    extras = [item for item in recovered if item["path"] not in matched_paths]
    return index, extras


def build_media_index():
    messages = read_json("messages.json")
    by_index = {m["index"]: m for m in messages}
    media_items = []
    for path in collect_files(MEDIA_DIR):
        rel = str(path.relative_to(ROOT)).replace("\\", "/")
        message_index = None
        match = re.search(r"_(\d+)_", path.name)
        if match:
            message_index = int(match.group(1))
        message = by_index.get(message_index) if message_index else None
        info = {
            "path": str(path),
            "relativePath": rel,
            "fileName": path.name,
            "extension": path.suffix.lower(),
            "bytes": path.stat().st_size,
            "sha256": sha256_file(path),
            "messageIndex": message_index,
            "time": message.get("time") if message else None,
            "sender": message.get("accountName") if message else None,
            "messageContent": message.get("content") if message else None,
        }
        try:
            with Image.open(path) as img:
                info.update({
                    "width": img.width,
                    "height": img.height,
                    "format": img.format,
                    "mode": img.mode,
                    "frames": getattr(img, "n_frames", 1),
                })
        except Exception as exc:
            info["imageError"] = repr(exc)
        media_items.append(info)
    return media_items


def extract_all():
    EXTRACT_DIR.mkdir(parents=True, exist_ok=True)
    MEDIA_EXTRACT_DIR.mkdir(parents=True, exist_ok=True)
    index, extras = build_attachment_index()
    extracted = []
    for row in index:
        recovered = row.get("recovered")
        if not recovered:
            extracted.append({**row, "extraction": None})
            continue
        path = Path(recovered["path"])
        out_dir = EXTRACT_DIR / f"{row['messageIndex']:03d}_{safe_stem(path.name)}"
        if out_dir.exists():
            shutil.rmtree(out_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        extraction = {"errors": [f"unsupported extension {path.suffix.lower()}"]}
        try:
            ext = path.suffix.lower()
            if ext == ".pdf":
                page_count, pdfinfo = pdf_page_count(path)
                extraction = extract_pdf(path, out_dir)
                if page_count is not None:
                    extraction["pdfInfoPageCount"] = page_count
                if pdfinfo:
                    extraction["pdfInfoTool"] = {
                        "ok": pdfinfo["ok"],
                        "returncode": pdfinfo["returncode"],
                        "tool": str(PDFINFO) if PDFINFO else None,
                    }
                if pdfinfo and pdfinfo["ok"]:
                    (out_dir / "pdfinfo.txt").write_text(pdfinfo["stdout"] + "\n" + pdfinfo["stderr"], encoding="utf-8")
                extraction["renderedSamples"] = render_pdf_samples(path, {"extraction": extraction})
            elif ext == ".docx":
                extraction = extract_docx(path, out_dir)
            elif ext == ".xlsx":
                extraction = extract_xlsx(path, out_dir)
            elif ext in {".md", ".txt"}:
                extraction = extract_md(path, out_dir)
        except Exception as exc:
            extraction = {"errors": [repr(exc)]}
        extracted.append({**row, "extractDir": str(out_dir), "extraction": extraction})

    media_index = build_media_index()
    write_json(ANALYSIS_DIR / "attachment_index.json", extracted)
    write_json(ANALYSIS_DIR / "attachment_extras.json", extras)
    write_json(ANALYSIS_DIR / "media_index.json", media_index)

    md = [
        "# Rokid 附件证据索引",
        "",
        f"- 群聊文件消息: {len(index)}",
        f"- 已恢复附件: {sum(1 for row in extracted if row.get('recovered'))}",
        f"- 未匹配额外附件: {len(extras)}",
        f"- 媒体文件: {len(media_index)}",
        "",
        "## 文件消息与恢复路径",
        "",
    ]
    for row in extracted:
        rec = row.get("recovered")
        ext = row.get("extraction") or {}
        md.extend([
            f"### #{row['messageIndex']} | {row['time']} | {row['sender']}",
            "",
            f"- 群聊文件名: {row['fileName']}",
            f"- 状态: {row['status']}",
            f"- 恢复路径: {rec['relativePath'] if rec else 'MISSING'}",
            f"- 字节数: {rec['bytes'] if rec else ''}",
            f"- 抽取目录: {row.get('extractDir', '')}",
            f"- 文本字数: {ext.get('textChars', '')}",
            f"- 页数/工作表/段落: {ext.get('pageCount', ext.get('sheetCount', ext.get('paragraphCount', '')))}",
            f"- 错误: {'; '.join(ext.get('errors', [])) if ext.get('errors') else '无'}",
            "",
        ])
    (ANALYSIS_DIR / "attachment_index.md").write_text("\n".join(md), encoding="utf-8")

    media_md = ["# Rokid 图片/表情媒体索引", ""]
    for item in media_index:
        media_md.append(
            f"- #{item.get('messageIndex') or ''} | {item.get('time') or ''} | "
            f"{item.get('sender') or ''} | {item['relativePath']} | "
            f"{item.get('width', '')}x{item.get('height', '')} | {item['bytes']} bytes"
        )
    (ANALYSIS_DIR / "media_index.md").write_text("\n".join(media_md) + "\n", encoding="utf-8")

    return {
        "attachmentMessages": len(index),
        "matched": sum(1 for row in extracted if row.get("recovered")),
        "missing": sum(1 for row in extracted if not row.get("recovered")),
        "extras": len(extras),
        "media": len(media_index),
        "textChars": sum((row.get("extraction") or {}).get("textChars", 0) for row in extracted),
    }


if __name__ == "__main__":
    summary = extract_all()
    print(json.dumps(summary, ensure_ascii=False, indent=2))
